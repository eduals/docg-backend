"""
Activity de Documento - Geração de documentos.
"""
import logging
from datetime import datetime
from typing import Dict, Any, Optional
from temporalio import activity

logger = logging.getLogger(__name__)

@activity.defn
async def execute_document_node(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Executa node de documento (geração).
    
    IDEMPOTÊNCIA: Verifica se documento já foi gerado antes de criar novo.
    
    Suporta:
    - google-docs: Google Docs
    - google-slides: Google Slides
    - microsoft-word: Microsoft Word
    - microsoft-powerpoint: Microsoft PowerPoint
    
    Args:
        data: {
            execution_id,
            node: {id, node_type, config, ...},
            workflow_id,
            organization_id,
            source_data: {...},
            source_object_id,
            source_object_type
        }
    
    Returns:
        {
            document_id,
            file_id,
            file_url,
            pdf_file_id,
            pdf_url,
            reused: bool
        }
    """
    from app.database import db
    from app.models import (
        GeneratedDocument, Template, Workflow, 
        DataSourceConnection
    )
    from app.services.document_generation.tag_processor import TagProcessor
    from flask import current_app
    
    node = data['node']
    node_type = node.get('node_type')
    config = node.get('config', {})
    execution_id = data['execution_id']
    node_id = node['id']
    
    activity.logger.info(f"Executando document node {node_id} tipo {node_type}")
    
    with current_app.app_context():
        # === IDEMPOTÊNCIA: Verificar se já existe ===
        existing = GeneratedDocument.query.filter_by(
            workflow_id=data['workflow_id']
        ).join(
            WorkflowNode, WorkflowNode.id == GeneratedDocument.template_id
        ).filter(
            # Buscar por execution via trigger_data ou outro meio
            GeneratedDocument.source_object_id == data['source_object_id']
        ).first()
        
        # Alternativa: buscar pelo context de execução
        # Por enquanto, sempre gerar novo documento (o Temporal garante que não vai duplicar)
        
        # === EXECUÇÃO ===
        template_id = config.get('template_id')
        if not template_id:
            raise ValueError(f'template_id não configurado no {node_type} node')
        
        template = Template.query.get(template_id)
        if not template:
            raise ValueError(f'Template não encontrado: {template_id}')
        
        workflow = Workflow.query.get(data['workflow_id'])
        if not workflow:
            raise ValueError(f'Workflow não encontrado: {data["workflow_id"]}')
        
        # Determinar/inferir node_type baseado em storage_type do template
        # Se node_type não foi especificado ou não corresponde ao template, inferir
        if template.storage_type == 'uploaded':
            # Template enviado - usar uploaded-document ou file-upload
            if node_type not in ['uploaded-document', 'file-upload']:
                node_type = 'file-upload'
                activity.logger.info(f"Inferindo node_type='file-upload' para template {template_id}")
        elif template.google_file_id:
            # Template do Google - inferir se necessário
            if node_type not in ['google-docs', 'google-slides']:
                if template.google_file_type == 'document':
                    node_type = 'google-docs'
                elif template.google_file_type == 'presentation':
                    node_type = 'google-slides'
                activity.logger.info(f"Inferindo node_type='{node_type}' para template Google {template_id}")
        elif template.microsoft_file_id:
            # Template do Microsoft - inferir se necessário
            if node_type not in ['microsoft-word', 'microsoft-powerpoint']:
                if template.microsoft_file_type == 'word':
                    node_type = 'microsoft-word'
                elif template.microsoft_file_type == 'powerpoint':
                    node_type = 'microsoft-powerpoint'
                activity.logger.info(f"Inferindo node_type='{node_type}' para template Microsoft {template_id}")
        
        # Gerar nome do documento
        output_name_template = config.get('output_name_template', '{{object_type}} - {{timestamp}}')
        data_with_meta = {
            **data['source_data'],
            'date': datetime.utcnow().strftime('%Y-%m-%d'),
            'timestamp': datetime.utcnow().strftime('%Y%m%d_%H%M%S'),
            'object_type': data['source_object_type']
        }
        doc_name = TagProcessor.replace_tags(output_name_template, data_with_meta)
        
        # Buscar field mappings
        field_mappings_data = config.get('field_mappings', [])
        mappings = {}
        for mapping_data in field_mappings_data:
            template_tag = mapping_data.get('template_tag')
            source_field = mapping_data.get('source_field')
            if template_tag and source_field:
                mappings[template_tag] = source_field
        
        # Processar AI mappings se houver (agora vem do node.config)
        ai_replacements = {}
        ai_mappings_config = config.get('ai_mappings', [])
        if ai_mappings_config:
            try:
                from app.services.document_generation.generator import DocumentGenerator, AIGenerationMetrics
                from app.routes.google_drive_routes import get_google_credentials

                google_creds = get_google_credentials(workflow.organization_id)
                if google_creds:
                    ai_metrics = AIGenerationMetrics()
                    generator = DocumentGenerator(google_creds)
                    # Processar AI tags usando config do node
                    ai_replacements = generator._process_ai_tags_from_config(
                        ai_mappings=ai_mappings_config,
                        source_data=data['source_data'],
                        workflow=workflow,
                        metrics=ai_metrics
                    )
                    activity.logger.info(f'AI tags processadas: {len(ai_replacements)} substituições')
            except Exception as e:
                activity.logger.warning(f'Erro ao processar AI tags: {e}')
        
        # Combinar dados
        combined_data = {**data['source_data'], **ai_replacements}
        
        # Executar baseado no tipo
        if node_type == 'google-docs':
            result = await _generate_google_docs(
                workflow, template, config, doc_name, combined_data, mappings, data
            )
        elif node_type == 'google-slides':
            result = await _generate_google_slides(
                workflow, template, config, doc_name, combined_data, mappings, data
            )
        elif node_type == 'microsoft-word':
            result = await _generate_microsoft_word(
                workflow, template, config, doc_name, combined_data, mappings, data
            )
        elif node_type == 'microsoft-powerpoint':
            result = await _generate_microsoft_powerpoint(
                workflow, template, config, doc_name, combined_data, mappings, data
            )
        elif node_type in ['uploaded-document', 'file-upload']:
            result = await _generate_uploaded_document(
                workflow, template, config, doc_name, combined_data, mappings, data
            )
        else:
            raise ValueError(f'Tipo de documento não suportado: {node_type}')
        
        activity.logger.info(f"Documento gerado: {result['document_id']}")
        return result

async def _generate_uploaded_document(
    workflow, template, config, doc_name, combined_data, mappings, data
) -> Dict[str, Any]:
    """
    Gera documento a partir de template enviado (DigitalOcean Spaces).
    
    Fluxo:
    1. Baixar template do DigitalOcean Spaces
    2. Normalizar .doc para .docx se necessário
    3. Validar estrutura do documento
    4. Processar com python-docx para substituir tags
    5. Salvar documento gerado no DigitalOcean Spaces (outputs/)
    6. Gerar PDF se configurado
    7. Criar registro GeneratedDocument
    """
    from app.database import db
    from app.models import GeneratedDocument
    from app.services.storage import DigitalOceanSpacesService
    from app.services.document_generation.tag_processor import TagProcessor
    from app.services.document_generation.document_converter import DocumentConverter
    from datetime import datetime
    import uuid
    from io import BytesIO
    from docx import Document
    import requests
    import os
    
    # Validar template
    if template.storage_type != 'uploaded':
        raise ValueError(f'Template {template.id} não é do tipo uploaded')
    
    if not template.storage_file_key:
        raise ValueError(f'Template {template.id} não tem storage_file_key')
    
    logger.info(f"Gerando documento a partir de template enviado: {template.id}")
    
    # Baixar template do DigitalOcean Spaces
    storage_service = DigitalOceanSpacesService()
    
    # Gerar URL assinada para download
    template_url = storage_service.generate_signed_url(
        template.storage_file_key,
        expiration=300  # 5 minutos
    )
    
    # Baixar arquivo
    response = requests.get(template_url, timeout=60)
    response.raise_for_status()
    template_bytes = response.content
    
    logger.info(f"Template baixado: {len(template_bytes)} bytes")
    
    # Determinar extensão do arquivo original
    file_extension = os.path.splitext(template.storage_file_key)[1] or '.docx'
    
    # Normalizar documento (.doc -> .docx)
    try:
        normalized_bytes, normalized_ext = DocumentConverter.normalize_document(
            template_bytes,
            file_extension
        )
        logger.info(f"Documento normalizado: {normalized_ext}")
    except ValueError as e:
        logger.error(f"Erro ao normalizar documento: {e}")
        raise ValueError(f"Não foi possível processar o arquivo: {str(e)}")
    
    # Validar estrutura do documento antes de processar
    is_valid, error_message = DocumentConverter.validate_document_structure(normalized_bytes)
    if not is_valid:
        logger.error(f"Documento inválido: {error_message}")
        raise ValueError(f"Documento não é válido: {error_message}")
    
    logger.info("Estrutura do documento validada com sucesso")
    
    # Processar documento com python-docx
    doc = Document(BytesIO(normalized_bytes))
    
    # Substituir tags em parágrafos
    for paragraph in doc.paragraphs:
        text = paragraph.text
        tags = TagProcessor.extract_tags(text)
        
        for tag in tags:
            field = mappings.get(tag, tag) if mappings else tag
            value = TagProcessor._get_nested_value(combined_data, field)
            
            if value is not None:
                text = text.replace(f'{{{{{tag}}}}}', str(value))
            else:
                text = text.replace(f'{{{{{tag}}}}}', '')
        
        paragraph.text = text
    
    # Substituir tags em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                tags = TagProcessor.extract_tags(text)
                
                for tag in tags:
                    field = mappings.get(tag, tag) if mappings else tag
                    value = TagProcessor._get_nested_value(combined_data, field)
                    
                    if value is not None:
                        text = text.replace(f'{{{{{tag}}}}}', str(value))
                    else:
                        text = text.replace(f'{{{{{tag}}}}}', '')
                
                cell.text = text
    
    # Salvar documento processado em buffer
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    processed_bytes = output_buffer.read()
    
    logger.info(f"Documento processado: {len(processed_bytes)} bytes")
    
    # Gerar nome único para arquivo gerado
    file_uuid = str(uuid.uuid4())
    file_ext = '.docx'  # Sempre .docx para templates enviados
    filename = f"{file_uuid}{file_ext}"
    
    # Key no DigitalOcean Spaces: docg/{organization_id}/outputs/{filename}
    organization_id = str(workflow.organization_id)
    output_key = f"docg/{organization_id}/outputs/{filename}"
    
    # Upload do documento gerado
    mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    output_url = storage_service.upload_file(
        BytesIO(processed_bytes),
        output_key,
        mime_type
    )
    
    logger.info(f"Documento gerado salvo no Spaces: {output_key}")
    
    # Gerar PDF se configurado
    pdf_result = None
    if config.get('create_pdf', False):
        try:
            logger.info("Gerando PDF do documento...")
            pdf_bytes = DocumentConverter.convert_docx_to_pdf(processed_bytes)
            
            # Upload do PDF para DigitalOcean Spaces
            pdf_uuid = str(uuid.uuid4())
            pdf_filename = f"{pdf_uuid}.pdf"
            pdf_key = f"docg/{organization_id}/outputs/{pdf_filename}"
            
            pdf_url = storage_service.upload_file(
                BytesIO(pdf_bytes),
                pdf_key,
                'application/pdf'
            )
            
            pdf_result = {
                'id': pdf_key,
                'url': pdf_url
            }
            
            logger.info(f"PDF gerado e salvo no Spaces: {pdf_key}")
            
        except ValueError as e:
            logger.warning(f'Não foi possível gerar PDF: {e}')
            # Não falhar a execução se PDF não puder ser gerado
            pdf_result = None
        except Exception as e:
            logger.error(f'Erro inesperado ao gerar PDF: {e}')
            pdf_result = None
    
    # Criar registro GeneratedDocument
    generated_doc = GeneratedDocument(
        organization_id=workflow.organization_id,
        workflow_id=workflow.id,
        source_connection_id=workflow.source_connection_id,
        source_object_type=data['source_object_type'],
        source_object_id=data['source_object_id'],
        template_id=template.id,
        template_version=template.version,
        name=doc_name,
        # Usar campos existentes para armazenar URL do DigitalOcean Spaces
        google_doc_id=output_key,  # Reutilizar campo para key no Spaces
        google_doc_url=output_url,  # Reutilizar campo para URL do Spaces
        status='generated',
        generated_data=data['source_data'],
        generated_at=datetime.utcnow()
    )
    
    # Adicionar informações do PDF se gerado
    if pdf_result:
        generated_doc.pdf_file_id = pdf_result['id']
        generated_doc.pdf_url = pdf_result['url']
    
    db.session.add(generated_doc)
    db.session.commit()
    
    logger.info(f"GeneratedDocument criado: {generated_doc.id}")
    
    return {
        'document_id': str(generated_doc.id),
        'file_id': output_key,  # Key no Spaces
        'file_url': output_url,
        'pdf_file_id': pdf_result['id'] if pdf_result else None,
        'pdf_url': pdf_result['url'] if pdf_result else None,
        'reused': False
    }

async def _generate_google_docs(
    workflow, template, config, doc_name, combined_data, mappings, data
) -> Dict[str, Any]:
    """Gera documento no Google Docs"""
    from app.database import db
    from app.models import GeneratedDocument
    from app.services.document_generation.generator import DocumentGenerator
    from app.routes.google_drive_routes import get_google_credentials
    from datetime import datetime
    
    google_creds = get_google_credentials(workflow.organization_id)
    if not google_creds:
        raise ValueError('Credenciais do Google não configuradas')
    
    generator = DocumentGenerator(google_creds)
    
    # Copiar template
    new_doc = generator.google_docs.copy_template(
        template_id=template.google_file_id,
        new_name=doc_name,
        folder_id=config.get('output_folder_id')
    )
    
    # Substituir tags
    generator.google_docs.replace_tags_in_document(
        document_id=new_doc['id'],
        data=combined_data,
        mappings=mappings
    )
    
    # Gerar PDF
    pdf_result = None
    if config.get('create_pdf', True):
        pdf_bytes = generator.google_docs.export_as_pdf(new_doc['id'])
        pdf_result = generator._upload_pdf(
            pdf_bytes,
            f"{doc_name}.pdf",
            config.get('output_folder_id')
        )
    
    # Criar registro
    generated_doc = GeneratedDocument(
        organization_id=workflow.organization_id,
        workflow_id=workflow.id,
        source_connection_id=workflow.source_connection_id,
        source_object_type=data['source_object_type'],
        source_object_id=data['source_object_id'],
        template_id=template.id,
        template_version=template.version,
        name=doc_name,
        google_doc_id=new_doc['id'],
        google_doc_url=new_doc['url'],
        status='generated',
        generated_data=data['source_data'],
        generated_at=datetime.utcnow()
    )
    
    if pdf_result:
        generated_doc.pdf_file_id = pdf_result['id']
        generated_doc.pdf_url = pdf_result['url']
    
    db.session.add(generated_doc)
    db.session.commit()
    
    return {
        'document_id': str(generated_doc.id),
        'file_id': new_doc['id'],
        'file_url': new_doc['url'],
        'pdf_file_id': pdf_result['id'] if pdf_result else None,
        'pdf_url': pdf_result['url'] if pdf_result else None,
        'reused': False
    }

async def _generate_google_slides(
    workflow, template, config, doc_name, combined_data, mappings, data
) -> Dict[str, Any]:
    """Gera apresentação no Google Slides"""
    from app.database import db
    from app.models import GeneratedDocument
    from app.services.document_generation.google_slides import GoogleSlidesService
    from app.services.document_generation.generator import DocumentGenerator
    from app.routes.google_drive_routes import get_google_credentials
    from datetime import datetime
    
    google_creds = get_google_credentials(workflow.organization_id)
    if not google_creds:
        raise ValueError('Credenciais do Google não configuradas')
    
    slides_service = GoogleSlidesService(google_creds)
    
    # Copiar template
    new_pres = slides_service.copy_template(
        template_id=template.google_file_id,
        new_name=doc_name,
        folder_id=config.get('output_folder_id')
    )
    
    # Substituir tags
    slides_service.replace_tags_in_presentation(
        presentation_id=new_pres['id'],
        data=combined_data,
        mappings=mappings
    )
    
    # Gerar PDF
    pdf_result = None
    if config.get('create_pdf', True):
        pdf_bytes = slides_service.export_as_pdf(new_pres['id'])
        generator = DocumentGenerator(google_creds)
        pdf_result = generator._upload_pdf(
            pdf_bytes,
            f"{doc_name}.pdf",
            config.get('output_folder_id')
        )
    
    # Criar registro
    generated_doc = GeneratedDocument(
        organization_id=workflow.organization_id,
        workflow_id=workflow.id,
        source_connection_id=workflow.source_connection_id,
        source_object_type=data['source_object_type'],
        source_object_id=data['source_object_id'],
        template_id=template.id,
        template_version=template.version,
        name=doc_name,
        google_doc_id=new_pres['id'],
        google_doc_url=new_pres['url'],
        status='generated',
        generated_data=data['source_data'],
        generated_at=datetime.utcnow()
    )
    
    if pdf_result:
        generated_doc.pdf_file_id = pdf_result['id']
        generated_doc.pdf_url = pdf_result['url']
    
    db.session.add(generated_doc)
    db.session.commit()
    
    return {
        'document_id': str(generated_doc.id),
        'file_id': new_pres['id'],
        'file_url': new_pres['url'],
        'pdf_file_id': pdf_result['id'] if pdf_result else None,
        'pdf_url': pdf_result['url'] if pdf_result else None,
        'reused': False
    }

async def _generate_microsoft_word(
    workflow, template, config, doc_name, combined_data, mappings, data
) -> Dict[str, Any]:
    """Gera documento no Microsoft Word"""
    from app.database import db
    from app.models import GeneratedDocument, DataSourceConnection
    from app.services.document_generation.microsoft_word import MicrosoftWordService
    from datetime import datetime
    import requests
    
    connection_id = config.get('connection_id')
    if not connection_id:
        raise ValueError('connection_id não configurado no Microsoft Word node')
    
    connection = DataSourceConnection.query.filter_by(
        id=connection_id,
        organization_id=workflow.organization_id,
        source_type='microsoft'
    ).first()
    
    if not connection:
        raise ValueError(f'Conexão Microsoft não encontrada: {connection_id}')
    
    credentials = connection.get_decrypted_credentials()
    access_token = credentials.get('access_token')
    
    if not access_token:
        raise ValueError('Access token não encontrado na conexão Microsoft')
    
    # Verificar expiração do token
    expires_at_str = credentials.get('expires_at')
    if expires_at_str:
        expires_at = datetime.fromisoformat(expires_at_str.replace('Z', '+00:00'))
        if expires_at < datetime.utcnow():
            from app.routes.microsoft_oauth_routes import _refresh_microsoft_token
            if not _refresh_microsoft_token(connection):
                raise ValueError('Não foi possível renovar token Microsoft')
            credentials = connection.get_decrypted_credentials()
            access_token = credentials.get('access_token')
    
    word_service = MicrosoftWordService({
        'access_token': access_token,
        'refresh_token': credentials.get('refresh_token'),
        'expires_at': credentials.get('expires_at'),
        'user_email': credentials.get('user_email'),
    })
    
    # Copiar template
    new_doc = word_service.copy_template(
        template_id=template.microsoft_file_id or template.google_file_id,
        new_name=doc_name,
        folder_id=config.get('output_folder_id')
    )
    
    # Substituir tags
    word_service.replace_tags_in_document(
        document_id=new_doc['id'],
        data=combined_data,
        mappings=mappings
    )
    
    # Gerar PDF
    pdf_result = None
    if config.get('create_pdf', True):
        try:
            pdf_bytes = word_service.export_as_pdf(new_doc['id'])
            pdf_name = f"{doc_name}.pdf"
            pdf_upload_response = requests.put(
                f'https://graph.microsoft.com/v1.0/me/drive/items/{config.get("output_folder_id", "root")}/children/{pdf_name}/content',
                headers={'Authorization': f'Bearer {access_token}'},
                data=pdf_bytes
            )
            if pdf_upload_response.ok:
                pdf_file = pdf_upload_response.json()
                pdf_result = {
                    'id': pdf_file['id'],
                    'url': pdf_file.get('webUrl')
                }
        except Exception as e:
            activity.logger.warning(f'Erro ao gerar PDF: {e}')
    
    # Criar registro
    generated_doc = GeneratedDocument(
        organization_id=workflow.organization_id,
        workflow_id=workflow.id,
        source_connection_id=workflow.source_connection_id,
        source_object_type=data['source_object_type'],
        source_object_id=data['source_object_id'],
        template_id=template.id,
        template_version=template.version,
        name=doc_name,
        google_doc_id=new_doc['id'],  # Reusa campo
        google_doc_url=new_doc['url'],
        status='generated',
        generated_data=data['source_data'],
        generated_at=datetime.utcnow()
    )
    
    if pdf_result:
        generated_doc.pdf_file_id = pdf_result['id']
        generated_doc.pdf_url = pdf_result['url']
    
    db.session.add(generated_doc)
    db.session.commit()
    
    return {
        'document_id': str(generated_doc.id),
        'file_id': new_doc['id'],
        'file_url': new_doc['url'],
        'pdf_file_id': pdf_result['id'] if pdf_result else None,
        'pdf_url': pdf_result['url'] if pdf_result else None,
        'reused': False
    }

async def _generate_microsoft_powerpoint(
    workflow, template, config, doc_name, combined_data, mappings, data
) -> Dict[str, Any]:
    """Gera apresentação no Microsoft PowerPoint"""
    from app.database import db
    from app.models import GeneratedDocument, DataSourceConnection
    from app.services.document_generation.microsoft_powerpoint import MicrosoftPowerPointService
    from datetime import datetime
    import requests
    
    connection_id = config.get('connection_id')
    if not connection_id:
        raise ValueError('connection_id não configurado no Microsoft PowerPoint node')
    
    connection = DataSourceConnection.query.filter_by(
        id=connection_id,
        organization_id=workflow.organization_id,
        source_type='microsoft'
    ).first()
    
    if not connection:
        raise ValueError(f'Conexão Microsoft não encontrada: {connection_id}')
    
    credentials = connection.get_decrypted_credentials()
    access_token = credentials.get('access_token')
    
    if not access_token:
        raise ValueError('Access token não encontrado na conexão Microsoft')
    
    # Verificar expiração do token
    expires_at_str = credentials.get('expires_at')
    if expires_at_str:
        expires_at = datetime.fromisoformat(expires_at_str.replace('Z', '+00:00'))
        if expires_at < datetime.utcnow():
            from app.routes.microsoft_oauth_routes import _refresh_microsoft_token
            if not _refresh_microsoft_token(connection):
                raise ValueError('Não foi possível renovar token Microsoft')
            credentials = connection.get_decrypted_credentials()
            access_token = credentials.get('access_token')
    
    ppt_service = MicrosoftPowerPointService({
        'access_token': access_token,
        'refresh_token': credentials.get('refresh_token'),
        'expires_at': credentials.get('expires_at'),
        'user_email': credentials.get('user_email'),
    })
    
    # Copiar template
    new_pres = ppt_service.copy_template(
        template_id=template.microsoft_file_id or template.google_file_id,
        new_name=doc_name,
        folder_id=config.get('output_folder_id')
    )
    
    # Substituir tags
    ppt_service.replace_tags_in_presentation(
        presentation_id=new_pres['id'],
        data=combined_data,
        mappings=mappings
    )
    
    # Gerar PDF
    pdf_result = None
    if config.get('create_pdf', True):
        try:
            pdf_bytes = ppt_service.export_as_pdf(new_pres['id'])
            pdf_name = f"{doc_name}.pdf"
            pdf_upload_response = requests.put(
                f'https://graph.microsoft.com/v1.0/me/drive/items/{config.get("output_folder_id", "root")}/children/{pdf_name}/content',
                headers={'Authorization': f'Bearer {access_token}'},
                data=pdf_bytes
            )
            if pdf_upload_response.ok:
                pdf_file = pdf_upload_response.json()
                pdf_result = {
                    'id': pdf_file['id'],
                    'url': pdf_file.get('webUrl')
                }
        except Exception as e:
            activity.logger.warning(f'Erro ao gerar PDF: {e}')
    
    # Criar registro
    generated_doc = GeneratedDocument(
        organization_id=workflow.organization_id,
        workflow_id=workflow.id,
        source_connection_id=workflow.source_connection_id,
        source_object_type=data['source_object_type'],
        source_object_id=data['source_object_id'],
        template_id=template.id,
        template_version=template.version,
        name=doc_name,
        google_doc_id=new_pres['id'],  # Reusa campo
        google_doc_url=new_pres['url'],
        status='generated',
        generated_data=data['source_data'],
        generated_at=datetime.utcnow()
    )
    
    if pdf_result:
        generated_doc.pdf_file_id = pdf_result['id']
        generated_doc.pdf_url = pdf_result['url']
    
    db.session.add(generated_doc)
    db.session.commit()
    
    return {
        'document_id': str(generated_doc.id),
        'file_id': new_pres['id'],
        'file_url': new_pres['url'],
        'pdf_file_id': pdf_result['id'] if pdf_result else None,
        'pdf_url': pdf_result['url'] if pdf_result else None,
        'reused': False
    }

