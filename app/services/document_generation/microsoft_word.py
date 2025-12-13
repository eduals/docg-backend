"""
Serviço para manipulação de Microsoft Word via Graph API.
Similar ao GoogleDocsService, mas para Word/OneDrive.
"""
from typing import Dict, Any, Optional
import requests
import logging
from .tag_processor import TagProcessor

logger = logging.getLogger(__name__)


class MicrosoftWordService:
    """
    Serviço para manipulação de Microsoft Word via Graph API.
    """
    
    def __init__(self, credentials: Dict[str, Any]):
        """
        Args:
            credentials: Dict com credenciais Microsoft (access_token, refresh_token, etc.)
                        ou string com access_token (compatibilidade)
        """
        if isinstance(credentials, str):
            # Compatibilidade: aceitar string como access_token
            self.access_token = credentials
        else:
            self.access_token = credentials.get('access_token')
        
        if not self.access_token:
            raise ValueError('access_token não fornecido')
        
        self.base_url = 'https://graph.microsoft.com/v1.0'
        self.headers = {
            'Authorization': f'Bearer {self.access_token}',
            'Content-Type': 'application/json'
        }
    
    def copy_template(self, template_id: str, new_name: str, folder_id: str = None) -> Dict:
        """
        Copia um template do Word no OneDrive/SharePoint.
        
        Args:
            template_id: ID do arquivo template (driveItem id)
            new_name: Nome do novo documento
            folder_id: ID da pasta de destino (opcional, usa mesma pasta do template se não fornecido)
        
        Returns:
            Dict com id e url do novo documento
        """
        # Obter informações do arquivo original
        file_info = requests.get(
            f'{self.base_url}/me/drive/items/{template_id}',
            headers=self.headers
        ).json()
        
        # Determinar pasta de destino
        parent_reference = file_info.get('parentReference', {})
        if folder_id:
            parent_id = folder_id
        else:
            parent_id = parent_reference.get('id')
        
        # Copiar arquivo
        copy_body = {
            'name': new_name
        }
        
        copy_response = requests.post(
            f'{self.base_url}/me/drive/items/{template_id}/copy',
            headers=self.headers,
            json=copy_body
        )
        
        if copy_response.status_code == 202:
            # Copy é assíncrono, obter o novo arquivo pelo nome
            # Aguardar um pouco e buscar
            import time
            time.sleep(1)
            
            # Buscar arquivo na pasta
            folder_items = requests.get(
                f'{self.base_url}/me/drive/items/{parent_id}/children',
                headers=self.headers,
                params={'$filter': f"name eq '{new_name}'"}
            ).json()
            
            if folder_items.get('value'):
                new_file = folder_items['value'][0]
                return {
                    'id': new_file['id'],
                    'url': new_file.get('webUrl', f"https://graph.microsoft.com/v1.0/me/drive/items/{new_file['id']}")
                }
        
        copy_response.raise_for_status()
        # Se não retornou 202, tentar obter do Location header
        location = copy_response.headers.get('Location')
        if location:
            # Fazer polling até o arquivo estar pronto
            import time
            for _ in range(10):
                time.sleep(1)
                status_response = requests.get(location, headers=self.headers)
                if status_response.status_code == 200:
                    status_data = status_response.json()
                    if status_data.get('status') == 'completed':
                        resource_id = status_data.get('resourceId', '').split('!')[-1]
                        return {
                            'id': resource_id,
                            'url': f"https://graph.microsoft.com/v1.0/me/drive/items/{resource_id}"
                        }
        
        raise Exception('Falha ao copiar arquivo Word')
    
    def get_document_content(self, document_id: str) -> bytes:
        """
        Obtém o conteúdo do documento Word como bytes.
        
        Args:
            document_id: ID do arquivo
        
        Returns:
            Bytes do arquivo Word (.docx)
        """
        response = requests.get(
            f'{self.base_url}/me/drive/items/{document_id}/content',
            headers={'Authorization': self.headers['Authorization']}
        )
        response.raise_for_status()
        return response.content
    
    def replace_tags_in_document(
        self,
        document_id: str,
        data: Dict[str, Any],
        mappings: Dict[str, str] = None
    ) -> None:
        """
        Substitui tags no documento Word.
        
        Nota: Microsoft Graph API não tem endpoint direto para editar conteúdo Word.
        Usaremos uma abordagem alternativa:
        1. Baixar o arquivo
        2. Usar python-docx para substituir tags
        3. Fazer upload do arquivo atualizado
        
        Args:
            document_id: ID do documento
            data: Dados para substituição
            mappings: Mapeamento de tags para campos
        """
        try:
            # Baixar arquivo
            docx_content = self.get_document_content(document_id)
            
            # Processar com python-docx
            from io import BytesIO
            from docx import Document
            
            doc = Document(BytesIO(docx_content))
            
            # Substituir tags em parágrafos
            for paragraph in doc.paragraphs:
                text = paragraph.text
                tags = TagProcessor.extract_tags(text)
                
                for tag in tags:
                    field = mappings.get(tag, tag) if mappings else tag
                    value = TagProcessor._get_nested_value(data, field)
                    
                    if value is not None:
                        text = text.replace(f'{{{{{tag}}}}}', str(value))
                    else:
                        text = text.replace(f'{{{{{tag}}}}}', '')
                
                paragraph.text = text
            
            # Substituir tags em tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text
                        tags = TagProcessor.extract_tags(text)
                        
                        for tag in tags:
                            field = mappings.get(tag, tag) if mappings else tag
                            value = TagProcessor._get_nested_value(data, field)
                            
                            if value is not None:
                                text = text.replace(f'{{{{{tag}}}}}', str(value))
                            else:
                                text = text.replace(f'{{{{{tag}}}}}', '')
                        
                        cell.text = text
            
            # Salvar em buffer
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            # Fazer upload do arquivo atualizado
            self._upload_document_content(document_id, output.read())
            
        except ImportError:
            logger.error('python-docx não instalado. Instale com: pip install python-docx')
            raise Exception('Biblioteca python-docx não disponível')
        except Exception as e:
            logger.exception(f'Erro ao substituir tags no Word: {str(e)}')
            raise
    
    def _upload_document_content(self, document_id: str, content: bytes) -> None:
        """
        Faz upload do conteúdo atualizado do documento.
        
        Args:
            document_id: ID do documento
            content: Conteúdo do arquivo (bytes)
        """
        # Microsoft Graph API requer upload em sessão para arquivos grandes
        # Para arquivos pequenos (< 4MB), podemos usar upload simples
        if len(content) < 4 * 1024 * 1024:  # 4MB
            response = requests.put(
                f'{self.base_url}/me/drive/items/{document_id}/content',
                headers={'Authorization': self.headers['Authorization']},
                data=content
            )
            response.raise_for_status()
        else:
            # Para arquivos grandes, usar upload session
            self._upload_large_file(document_id, content)
    
    def _upload_large_file(self, document_id: str, content: bytes) -> None:
        """
        Faz upload de arquivo grande usando upload session.
        """
        # Criar upload session
        session_response = requests.post(
            f'{self.base_url}/me/drive/items/{document_id}/createUploadSession',
            headers=self.headers,
            json={
                'item': {
                    '@microsoft.graph.conflictBehavior': 'replace'
                }
            }
        )
        session_response.raise_for_status()
        upload_url = session_response.json()['uploadUrl']
        
        # Fazer upload
        upload_response = requests.put(
            upload_url,
            headers={'Content-Length': str(len(content))},
            data=content
        )
        upload_response.raise_for_status()
    
    def export_as_pdf(self, document_id: str) -> bytes:
        """
        Exporta documento Word como PDF.
        
        Args:
            document_id: ID do documento
        
        Returns:
            Bytes do PDF
        """
        response = requests.get(
            f'{self.base_url}/me/drive/items/{document_id}/content',
            headers={
                'Authorization': self.headers['Authorization'],
                'Accept': 'application/pdf'
            }
        )
        response.raise_for_status()
        return response.content

