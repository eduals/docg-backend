"""
Serviço para conversão de documentos e validação.
Suporta:
- Conversão de .doc para .docx (usando LibreOffice)
- Conversão de .docx para PDF (usando LibreOffice)
- Validação de estrutura de documentos
"""
import logging
import subprocess
import tempfile
import os
from pathlib import Path
from io import BytesIO
from typing import Optional, Tuple
from docx import Document

logger = logging.getLogger(__name__)


class DocumentConverter:
    """Utilitário para conversão e validação de documentos"""
    
    @staticmethod
    def convert_doc_to_docx(doc_bytes: bytes) -> bytes:
        """
        Converte arquivo .doc (Word 97-2003) para .docx usando LibreOffice.
        
        Args:
            doc_bytes: Bytes do arquivo .doc
            
        Returns:
            Bytes do arquivo .docx convertido
            
        Raises:
            ValueError: Se LibreOffice não estiver disponível ou conversão falhar
        """
        # Verificar se LibreOffice está disponível
        try:
            result = subprocess.run(
                ['soffice', '--version'],
                capture_output=True,
                timeout=5
            )
            if result.returncode != 0:
                raise ValueError('LibreOffice não está disponível no sistema')
        except (FileNotFoundError, subprocess.TimeoutExpired):
            raise ValueError(
                'LibreOffice não está instalado. '
                'Instale LibreOffice para suportar arquivos .doc'
            )
        
        # Criar arquivo temporário para input
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as input_file:
            input_file.write(doc_bytes)
            input_path = input_file.name
        
        try:
            # Criar diretório temporário para output
            output_dir = tempfile.mkdtemp()
            
            # Converter usando LibreOffice
            result = subprocess.run(
                [
                    'soffice',
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', output_dir,
                    input_path
                ],
                capture_output=True,
                timeout=30
            )
            
            if result.returncode != 0:
                error_msg = result.stderr.decode('utf-8', errors='ignore')
                raise ValueError(f'Erro ao converter .doc para .docx: {error_msg}')
            
            # Encontrar arquivo convertido
            input_name = Path(input_path).stem
            output_path = os.path.join(output_dir, f'{input_name}.docx')
            
            if not os.path.exists(output_path):
                raise ValueError('Arquivo convertido não foi encontrado')
            
            # Ler arquivo convertido
            with open(output_path, 'rb') as f:
                docx_bytes = f.read()
            
            # Limpar arquivos temporários
            os.unlink(input_path)
            os.unlink(output_path)
            os.rmdir(output_dir)
            
            logger.info(f'Documento .doc convertido para .docx: {len(docx_bytes)} bytes')
            return docx_bytes
            
        except subprocess.TimeoutExpired:
            # Limpar em caso de timeout
            if os.path.exists(input_path):
                os.unlink(input_path)
            raise ValueError('Timeout ao converter documento .doc')
        except Exception as e:
            # Limpar em caso de erro
            if os.path.exists(input_path):
                os.unlink(input_path)
            raise ValueError(f'Erro ao converter .doc: {str(e)}')
    
    @staticmethod
    def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
        """
        Converte arquivo .docx para PDF usando LibreOffice.
        
        Args:
            docx_bytes: Bytes do arquivo .docx
            
        Returns:
            Bytes do arquivo PDF gerado
            
        Raises:
            ValueError: Se LibreOffice não estiver disponível ou conversão falhar
        """
        # Verificar se LibreOffice está disponível
        try:
            result = subprocess.run(
                ['soffice', '--version'],
                capture_output=True,
                timeout=5
            )
            if result.returncode != 0:
                raise ValueError('LibreOffice não está disponível no sistema')
        except (FileNotFoundError, subprocess.TimeoutExpired):
            raise ValueError(
                'LibreOffice não está instalado. '
                'Instale LibreOffice para gerar PDFs'
            )
        
        # Criar arquivo temporário para input
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as input_file:
            input_file.write(docx_bytes)
            input_path = input_file.name
        
        try:
            # Criar diretório temporário para output
            output_dir = tempfile.mkdtemp()
            
            # Converter usando LibreOffice
            result = subprocess.run(
                [
                    'soffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    input_path
                ],
                capture_output=True,
                timeout=60  # PDF pode demorar mais
            )
            
            if result.returncode != 0:
                error_msg = result.stderr.decode('utf-8', errors='ignore')
                raise ValueError(f'Erro ao converter .docx para PDF: {error_msg}')
            
            # Encontrar arquivo convertido
            input_name = Path(input_path).stem
            output_path = os.path.join(output_dir, f'{input_name}.pdf')
            
            if not os.path.exists(output_path):
                raise ValueError('Arquivo PDF não foi gerado')
            
            # Ler arquivo PDF gerado
            with open(output_path, 'rb') as f:
                pdf_bytes = f.read()
            
            # Limpar arquivos temporários
            os.unlink(input_path)
            os.unlink(output_path)
            os.rmdir(output_dir)
            
            logger.info(f'Documento .docx convertido para PDF: {len(pdf_bytes)} bytes')
            return pdf_bytes
            
        except subprocess.TimeoutExpired:
            # Limpar em caso de timeout
            if os.path.exists(input_path):
                os.unlink(input_path)
            raise ValueError('Timeout ao gerar PDF')
        except Exception as e:
            # Limpar em caso de erro
            if os.path.exists(input_path):
                os.unlink(input_path)
            raise ValueError(f'Erro ao gerar PDF: {str(e)}')
    
    @staticmethod
    def validate_document_structure(docx_bytes: bytes) -> Tuple[bool, Optional[str]]:
        """
        Valida estrutura básica de um documento .docx.
        
        Verifica:
        - Se o arquivo é um .docx válido
        - Se tem pelo menos um parágrafo ou tabela
        - Se não está corrompido
        
        Args:
            docx_bytes: Bytes do arquivo .docx
            
        Returns:
            Tuple (is_valid, error_message)
            - is_valid: True se documento é válido
            - error_message: Mensagem de erro se inválido, None se válido
        """
        try:
            # Tentar abrir documento
            doc = Document(BytesIO(docx_bytes))
            
            # Verificar se tem conteúdo
            has_paragraphs = len(doc.paragraphs) > 0
            has_tables = len(doc.tables) > 0
            
            if not has_paragraphs and not has_tables:
                return False, 'Documento não contém parágrafos ou tabelas'
            
            # Verificar se pelo menos um parágrafo ou célula tem texto
            has_text = False
            
            # Verificar parágrafos
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    has_text = True
                    break
            
            # Verificar tabelas
            if not has_text:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text and cell.text.strip():
                                has_text = True
                                break
                        if has_text:
                            break
                    if has_text:
                        break
            
            if not has_text:
                return False, 'Documento não contém texto'
            
            return True, None
            
        except Exception as e:
            return False, f'Erro ao validar documento: {str(e)}'
    
    @staticmethod
    def normalize_document(doc_bytes: bytes, file_extension: str) -> Tuple[bytes, str]:
        """
        Normaliza documento para .docx.
        
        Se for .doc, converte para .docx.
        Se já for .docx, retorna como está.
        
        Args:
            doc_bytes: Bytes do arquivo original
            file_extension: Extensão do arquivo (.doc ou .docx)
            
        Returns:
            Tuple (docx_bytes, normalized_extension)
        """
        file_extension = file_extension.lower()
        
        if file_extension == '.docx':
            return doc_bytes, '.docx'
        elif file_extension == '.doc':
            logger.info('Convertendo arquivo .doc para .docx')
            docx_bytes = DocumentConverter.convert_doc_to_docx(doc_bytes)
            return docx_bytes, '.docx'
        else:
            raise ValueError(f'Extensão não suportada: {file_extension}')
